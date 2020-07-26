from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx2pdf import convert

#Initialsing Document object
document = Document()

#Creating a header section
section = document.sections[0]
header = section.header
footer = section.footer

#Creating custom styles
styles = document.styles
heading = styles.add_style('Practical_Number',WD_STYLE_TYPE.PARAGRAPH)
sub_heading = styles.add_style('Sub_Head',WD_STYLE_TYPE.PARAGRAPH)
content_style = styles.add_style('Content',WD_STYLE_TYPE.PARAGRAPH)

#Formatting built-in Header style
header_style = document.styles['Header']
header_font = header_style.font
header_font.name = 'Times New Roman'
header_font.size = Pt(12)
header_font.bold =True

footer_style = document.styles['Footer']
footer_font = footer_style.font
footer_font.name = 'Times New Roman'
footer_font.size = Pt(12)
footer_font.bold = True

#Defining custom styles
head = document.styles['Practical_Number']
head_font = head.font
head_font.bold=True
head_font.name= 'Times New Roman'
head_font.size = Pt(16)

sub_head_style = document.styles['Sub_Head']
sub_head_font = sub_head_style.font
sub_head_font.name='Times New Roman'
sub_head_font.size= Pt(14)
sub_head_font.bold=True

content = document.styles['Content']
content_font = content.font
content_font.name='Times New Roman'
content_font.size = Pt(12)



#----------------------------------------

#Actual Content starts

#----------------------------------------

#Variables for Data

header_data="18DCS035\t\tCS347"     #Insert Header Here
footer_data = "DEPSTAR(CSE)"
title_data="Practical 1"            #Insert Heading Here
aim_data="To create a Python program to generate word files"  #Insert Aim Here
program_code_data="Code here"       #Insert Code Here
picture_path="C:/Users/Admin/Desktop/photo.jpg"
conclusion_data="Conclusion"        #Insert Conclusion Here


#Adding header
header_content = header.paragraphs[0]
header_content.text = header_data
header_content.style = document.styles["Header"]

footer_content = footer.paragraphs[0]
footer_content.text= footer_data
footer_content.style = document.styles["Footer"]

#Adding title
title = document.add_paragraph(title_data,style='Practical_Number')
title.alignment=WD_ALIGN_PARAGRAPH.CENTER

#Adding sub-title
aim = document.add_paragraph("Aim",style='Sub_Head')

#Adding content to sub-title
aim_content = document.add_paragraph(aim_data,style='Content')
aim_content.alignment= WD_ALIGN_PARAGRAPH.JUSTIFY
program_code = document.add_paragraph("Program Code",style='Sub_Head')


program_code_content = document.add_paragraph(program_code_data,style='Content')
program_code_content.alignment= WD_ALIGN_PARAGRAPH.JUSTIFY
output = document.add_paragraph('Output',style='Sub_Head') 
document.add_picture(picture_path,width=Inches(1.25))

picture = document.paragraphs[-1]
picture.alignment= WD_ALIGN_PARAGRAPH.CENTER
conclusion = document.add_paragraph('Conclusion',style='Sub_Head')
conclusion_content = document.add_paragraph(conclusion_data,style='Content')
conclusion_content.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

document.save('template_test.docx')

convert('template_test.docx')
